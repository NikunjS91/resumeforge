import re
import pdfplumber
from docx import Document as DocxDocument
from collections import defaultdict


# ─── PDF EXTRACTION ──────────────────────────────────────────────────────────

def extract_pdf(file_path: str) -> dict:
    """
    Layout-aware PDF extraction using bounding box word positions.
    Reconstructs lines by grouping words with similar Y coordinates,
    which prevents multi-column skill tables from bleeding into each other.
    """
    all_lines = []
    page_count = 0

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False,
                use_text_flow=False,
            )

            if not words:
                # Fallback to basic extract_text for scanned/image PDFs
                text = page.extract_text()
                if text:
                    all_lines.extend(text.split("\n"))
                continue

            # Group words by their Y position (rounded to nearest 2px)
            lines_by_y = defaultdict(list)
            for word in words:
                y_key = round(word["top"] / 2) * 2
                lines_by_y[y_key].append(word)

            # Sort lines top to bottom, words left to right within each line
            for y_key in sorted(lines_by_y.keys()):
                line_words = sorted(lines_by_y[y_key], key=lambda w: w["x0"])
                line_text = " ".join(w["text"] for w in line_words).strip()
                if line_text:
                    all_lines.append(line_text)

            # Blank line between pages
            all_lines.append("")

    raw_text = "\n".join(all_lines).strip()
    # Rejoin words hyphenated across PDF lines (e.g. "CloudForma-\ntion" → "CloudFormation")
    raw_text = re.sub(r'-\n(\w)', r'\1', raw_text)
    return {"raw_text": raw_text, "page_count": page_count}


# ─── DOCX EXTRACTION ─────────────────────────────────────────────────────────

def extract_docx(file_path: str) -> dict:
    """Extract raw text from a DOCX file (paragraphs + table cells)."""
    doc = DocxDocument(file_path)
    lines = []

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())

    raw_text = "\n".join(lines)
    return {"raw_text": raw_text, "page_count": 1}


# ─── UNIFIED ENTRY POINT ─────────────────────────────────────────────────────

def extract(file_path: str, file_format: str) -> dict:
    """Route to correct extractor based on file format."""
    if file_format == "pdf":
        return extract_pdf(file_path)
    elif file_format == "docx":
        return extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported format: {file_format}")
